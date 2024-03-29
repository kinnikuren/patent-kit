# -*- coding: utf-8 -*-
"""
Created on Tue Oct 16 19:57:56 2018

@author: alanyliu
"""

import re
import pydot
import datetime
import docx
from docx.shared import Pt 

import file_reader

def add_original_identifier(claims_string_regex_list):
    list_with_original = []
    for i in range(len(claims_string_regex_list)):
        current_entry = claims_string_regex_list[i].strip()
        if i%2 != 0:
            current_entry += "\t(Original)  "
        list_with_original.append(current_entry)
    #print(list_with_original)   

    return list_with_original

def get_claims_string(filepath):
    claims_string = ""
    if (filepath.endswith(".txt") or filepath.endswith(".pdf")):    
        claims_string = file_reader.get_string_from_file(filepath).strip()
    elif (filepath.endswith(".docx")):
        claims_docx = docx.Document(filepath)
        style = claims_docx.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        for i in range(len(claims_docx.paragraphs)):
            p = claims_docx.paragraphs[i]
            
            #p.style = claims_docx.styles['Normal']
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            print(p.style.name)
            #print(i)
            print(p.text)
           
            claim_no_regex = r'([0-9]{1,2}\.)'
            #if(re.match(claim_no_regex,p.text)):
            line_split = re.split(claim_no_regex,p.text)
            if (len(line_split) > 1):
                print(line_split)
                line_split[1] += "\t(Original)  "
                claims_docx.paragraphs[i].text = line_split[1] + line_split[2].strip()
                print(claims_docx.paragraphs[i].text)
           
        #claims_docx.styles.style.__ParagraphStyle.font.size = Pt(14)
        claims_docx.save("output/test.docx")
        raise SystemExit("docx")
    else:
        print("Unsupported file format")
        raise SystemExit("Unsupported file format")

    return claims_string

def remove_tabs(claims_string):
    #remove initial tabs
    claims_string_linebreak_split = claims_string.split('\n')
    removed_tab_list = []
    for line in claims_string_linebreak_split:
        #print(line.strip())
        removed_tab_list.append(line.strip())
    new_claims_string = '\n'.join(removed_tab_list)
    print('tabs removed:')
    print(new_claims_string)   

    return new_claims_string

def create_claims_dict(claims_string):
    claims_string = remove_tabs(claims_string)                                                     
    claims_dict = {}

    claims_list = re.findall(r'[0-9]{1,2}\..*',claims_string,re.M)
    
    print('\nFINDALL results:')
    print(claims_list)
    print('\n\n')


    for i in claims_list:
        claimNo = re.search(r'[0-9]{1,2}',i).group(0)
        print('CLAIM ' + claimNo)
        #print(type(claimNo))
        #matchObj = re.search(r'claim [0-9]{1,2}',i,re.I)
        
        claims_dict[claimNo] = []
    
        #regex to find dependent claims
        depClaimRegex = r'(claim) (\d+)'
        
        matchObj = re.search(depClaimRegex,i,re.I)
        if matchObj:
            parClaimStr = matchObj.group(0)
            
            #depClaimNo = re.sub(depClaimRegex, r'\1', i)
            parClaimNo = parClaimStr.lower().replace('claim ','')
            print('parent: ' + parClaimNo)
                    
            if parClaimNo in claims_dict:
                claims_dict[parClaimNo].append(claimNo)
                print('added dependent claim {} to parent claim {}'.format(claimNo, parClaimNo))
            #depClaimNo = depClaimStr.replace()
            #print(matchObj.group(0))
                print(claims_dict)
            
        else:
            print('no match')
        
        
        #claims.update()
    
    return claims_dict

def generate_claim_chart(claims_dict, filepath="./"):
    graph = pydot.Dot(graph_type='graph')
    location = file_reader.get_folder(filepath)

    for i in claims_dict:
        #print(i)
        # we can get right into action by "drawing" edges between the nodes in our graph
        # we do not need to CREATE nodes, but if you want to give them some custom style
        # then I would recomend you to do so... let's cover that later
        # the pydot.Edge() constructor receives two parameters, a source node and a destination
        # node, they are just strings like you can see
        node = pydot.Node(i)
        graph.add_node(node)
        for j in claims_dict[i]:
            edge = pydot.Edge(i, j)
            graph.add_edge(edge)
        
        # and we obviosuly need to add the edge to our graph
    
    #currentTime = str(datetime.datetime.now()).split('.')[0].replace('-','').replace(' ','').replace(':','')
    
    #filepath = '{}_{}.png'.format(filepath.replace('.txt',''),currentTime)    
    #output_path = location + "claim_chart.png"
    output_path = location + "static/output/claim_chart.png"
    graph.write_png(output_path)
    print("flowchart generated at {}...".format(output_path))

def main():
    filepath=""
    #filepath = '8836S-1189 claims.txt'
    #filepath = "cases/8054L-1152/8054L-1152 claims.txt"
    
    #Temp
    #filepath = "input/8836S_1317_claims.txt"
    
    filepath = file_reader.get_filepath()
        
    #get claims as a string
    claims_string = get_claims_string(filepath)
           
    list2 = re.split(r'([0-9]{1,2}\.)',claims_string)
    
    #print(list)
    #print(len(list))
    
    #print(list2)
    
    list3 = add_original_identifier(list2)
    #print(list3)

    claims_dict = create_claims_dict(claims_string)
    print(claims_dict)
    
    generate_claim_chart(claims_dict, filepath)

#main()
