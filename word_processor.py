# -*- coding: utf-8 -*-
"""
Created on Tue Oct 16 22:25:02 2018

@author: alanyliu
"""

import docx

#def get_document(file_name):
    

def test_main():
    document = docx.Document('cases/8054S-1306/11-09-CY6010US-order.docx')
    count = 1
    for p in document.paragraphs:
        print(count)
        print(p.text)
        count += 1
        if (type(p) != str):
            #print(type(p))
            #print(p.text)
            pass
        if (type(p) == str):
            #print(type(p))
            print(p.text)
    
    for i, p in enumerate(document.paragraphs):
        if (p.text.isupper()):
            print(p.text)
            print(i)
        
    print(type(document.paragraphs))
    print(len(document.paragraphs))
    #print('\n'.join(document.paragraphs))
    
    print(document.sections)
    print(len(document.sections))
    for section in document.sections:
        print(section)
        
test_main()
    