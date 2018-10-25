# -*- coding: utf-8 -*-
"""
Created on Tue Oct 16 22:25:02 2018

@author: alanyliu
"""

from docx import Document


document = Document('8054S-1225-FINAL APPLN.docx')
count = 1
for p in document.paragraphs:
    #print(count)
    #print(p.text)
    count += 1
    if (type(p) != str):
        print(type(p))
        print(p.text)


print(type(document.paragraphs))
#print('\n'.join(document.paragraphs))